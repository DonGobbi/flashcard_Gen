from docx import Document
from docx.shared import Pt
import os

def create_sample_docx():
    doc = Document()
    
    # Add a title
    doc.add_heading('Mathematics Flashcards', 0)
    
    # Add some flashcard content
    flashcards = [
        ("What is the Pythagorean theorem?", "In a right triangle, a² + b² = c², where c is the hypotenuse"),
        ("What is a prime number?", "A number that has exactly two factors: 1 and itself"),
        ("What is the quadratic formula?", "x = (-b ± √(b² - 4ac)) / 2a"),
        ("What is a derivative?", "The rate of change of a function with respect to a variable"),
        ("What is integration?", "The process of finding the function whose derivative is the given function")
    ]
    
    for question, answer in flashcards:
        p = doc.add_paragraph()
        p.add_run(f"{question}: ").bold = True
        p.add_run(answer)
        doc.add_paragraph()  # Add space between flashcards
    
    # Save the document in the test_files directory
    output_path = os.path.join(os.path.dirname(__file__), 'mathematics.docx')
    doc.save(output_path)

if __name__ == '__main__':
    create_sample_docx()
